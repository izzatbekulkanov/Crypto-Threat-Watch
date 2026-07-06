"""Word (.docx) hisobot yaratuvchi modul — Professional va chiroyli hisobotlar uchun."""

import os
from datetime import datetime, timezone, timedelta
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Ranglar palitrasi
COLOR_PRIMARY = RGBColor(31, 78, 121)     # Deep Navy #1F4E79
COLOR_SECONDARY = RGBColor(128, 128, 128) # Grey
COLOR_TEXT = RGBColor(51, 51, 51)         # Charcoal
COLOR_WHITE = RGBColor(255, 255, 255)

# Risk ranglari
RISK_COLORS = {
    "LOW": "2ECC71",      # Green
    "MEDIUM": "F1C40F",   # Yellow
    "HIGH": "E67E22",     # Orange
    "CRITICAL": "E74C3C", # Red
}

REPORT_TEXTS = {
    "uz": {
        "title": "KENGAYTIRILGAN AUDIT HISOBOTI",
        "subtitle": "Crypto Threat Watch - Professional Monitoring Tizimi",
        "wallet_info": "Hamyon Ma'lumotlari",
        "network": "Tarmoq",
        "address": "Manzil",
        "risk_level": "Xavf Darajasi",
        "status": "Holati",
        "date": "Skanerlangan vaqt",
        "assets_summary": "Aktivlar Balansi va Aylanmasi (Ichki aylanmalarsiz)",
        "table_symbol": "Aktiv",
        "table_balance": "Hozirgi Balans",
        "table_income": "Kirim (IN)",
        "table_outcome": "Chiqim (OUT)",
        "table_net": "Sof Qoldiq (Net)",
        "table_volume": "Aylanma (Volume)",
        "table_txs": "Tranzaksiyalar",
        "swaps_title": "Ichki Almashtirishlar (DEX Swaps) - Turnoverdan chiqarilganlar",
        "swaps_intro": "Hamyonning o'zida amalga oshirilgan valyuta ayirboshlash (DEX swaps) operatsiyalari aniqlandi. Ular hamyonning tashqi kirim-chiqim aylanmalariga qo'shilmagan:",
        "no_swaps": "Ichki almashtirishlar aniqlanmadi.",
        "table_time": "Vaqt",
        "table_tx_hash": "Tranzaksiya Hash",
        "table_sent": "Yuborildi",
        "table_received": "Qabul qilindi",
        "yearly_title": "Yillik Faollik (Native Aktiv)",
        "yearly_year": "Yil",
        "risk_details_title": "Xavf Darajasi Tahlili",
        "risk_score_desc": "Hamyonning xulq-atvori va tranzaksiyalariga ko'ra baholangan risk ko'rsatkichlari.",
        "recommendations_title": "Xavfsizlik Tavsiyalari",
        "rec_low": "🟢 Hamyon past xavf guruhida. Odatiy xavfsizlik choralarini davom ettiring. Shubhali shartnomalar bilan o'zaro aloqada bo'lmang.",
        "rec_med": "🟡 Hamyon o'rta xavf guruhida. Tranzaksiyalar soni yoki ba'zi shubhali aylanmalar mavjud. Hamyon egalarining kimligini tekshirish tavsiya etiladi.",
        "rec_high": "🟠 Hamyon yuqori xavf guruhida! Katta hajmdagi aylanmalar yoki kirim/chiqim nomutanosibligi aniqlangan. Ushbu manzil bilan ishlashda ehtiyot bo'ling.",
        "rec_critical": "🔴 CRITICAL - Hamyon o'ta yuqori xavf ostida! Mumkin bo'lgan noqonuniy aylanmalar, mikserlar yoki shubhali faollik belgilari mavjud. Bloklash yoki to'liq monitoringga olish tavsiya etiladi.",
        "footer_text": "Crypto Threat Watch - Kiberxavfsizlik Bo'limi Maxsus Hisoboti.",
        
        # Yangi tranzaksiyalar jadvali uchun
        "tx_history_title": "Barcha Tranzaksiyalar Tarixi (Tashqi hamyonlar bilan)",
        "table_num": "№",
        "table_direction": "Turi",
        "table_counterparty": "Qarshi Hamyon (From/To)",
        "table_amount": "Qiymati",
        "table_asset": "Valyuta",
        "no_txs_text": "Tashqi hamyonlar bilan tranzaksiyalar aniqlanmadi."
    },
    "ru": {
        "title": "РАСШИРЕННЫЙ ОТЧЕТ ПО АУДИТУ",
        "subtitle": "Crypto Threat Watch - Профессиональная Система Мониторинга",
        "wallet_info": "Информация о Кошельке",
        "network": "Сеть",
        "address": "Адрес",
        "risk_level": "Уровень Риска",
        "status": "Статус",
        "date": "Время сканирования",
        "assets_summary": "Баланс Активов и Обороты (Исключая внутренние обмены)",
        "table_symbol": "Актив",
        "table_balance": "Текущий Баланс",
        "table_income": "Приход (IN)",
        "table_outcome": "Расход (OUT)",
        "table_net": "Чистый Остаток",
        "table_volume": "Оборот (Volume)",
        "table_txs": "Транзакции",
        "swaps_title": "Внутренние Обмены (DEX Swaps) - Исключены из оборота",
        "swaps_intro": "Обнаружены операции обмена валют (DEX swaps), проведенные внутри кошелька. Они исключены из внешнего прихода/расхода:",
        "no_swaps": "Внутренние обмены не обнаружены.",
        "table_time": "Время",
        "table_tx_hash": "Хэш Транзакции",
        "table_sent": "Отправлено",
        "table_received": "Получено",
        "yearly_title": "Ежегодная Активность (Native Актив)",
        "yearly_year": "Год",
        "risk_details_title": "Анализ Уровня Риска",
        "risk_score_desc": "Показатели риска, оцененные на основе активности и оборотов кошелька.",
        "recommendations_title": "Рекомендации по Безопасности",
        "rec_low": "🟢 Кошелек в группе низкого риска. Продолжайте соблюдать стандартные меры безопасности. Избегайте взаимодействия с подозрительными контрактами.",
        "rec_med": "🟡 Кошелек в группе среднего риска. Есть значительное число транзакций или некоторые дисбалансы. Рекомендуется проверить личность владельца.",
        "rec_high": "🟠 Кошелек в группе высокого риска! Обнаружены большие объемы или дисбаланс прихода/расхода. Соблюдайте осторожность при работе с этим адресом.",
        "rec_critical": "🔴 КРИТИЧЕСКИЙ - Кошелек под крайне высоким риском! Возможные незаконные обороты, использование миксеров или подозрительная активность. Рекомендуется заблокировать или взять на полный мониторинг.",
        "footer_text": "Crypto Threat Watch - Специальный Отчет Департамента Кибербезопасности.",
        
        # Yangi tranzaksiyalar jadvali uchun
        "tx_history_title": "История всех транзакций (С внешними кошельками)",
        "table_num": "№",
        "table_direction": "Тип",
        "table_counterparty": "Контрагент (From/To)",
        "table_amount": "Сумма",
        "table_asset": "Валюта",
        "no_txs_text": "Транзакции с внешними кошельками не найдены."
    },
    "en": {
        "title": "EXTENDED AUDIT REPORT",
        "subtitle": "Crypto Threat Watch - Professional Monitoring System",
        "wallet_info": "Wallet Information",
        "network": "Network",
        "address": "Address",
        "risk_level": "Risk Level",
        "status": "Status",
        "date": "Audit Timestamp",
        "assets_summary": "Asset Balances & Volume Summary (Excluding Swaps)",
        "table_symbol": "Asset",
        "table_balance": "Current Balance",
        "table_income": "Income (IN)",
        "table_outcome": "Outcome (OUT)",
        "table_net": "Net Balance",
        "table_volume": "Volume",
        "table_txs": "Tx Count",
        "swaps_title": "Internal Swaps (DEX Exchanges) - Excluded from turnovers",
        "swaps_intro": "DEX swap transactions within the same wallet detected. They have been excluded from the external incoming/outgoing volume:",
        "no_swaps": "No internal swaps detected.",
        "table_time": "Time",
        "table_tx_hash": "Transaction Hash",
        "table_sent": "Sent",
        "table_received": "Received",
        "yearly_title": "Yearly Activity (Native Asset)",
        "yearly_year": "Year",
        "risk_details_title": "Risk Assessment Analysis",
        "risk_score_desc": "Risk indicators evaluated based on the wallet's behavior and transactions.",
        "recommendations_title": "Security Recommendations",
        "rec_low": "🟢 Wallet is in the low-risk category. Continue standard security practices. Avoid interacting with unverified contracts.",
        "rec_med": "🟡 Wallet is in the medium-risk category. Moderate transaction count or volume imbalance detected. Verify counterparty identity.",
        "rec_high": "🟠 Wallet is in the high-risk category! High volumes or significant inflow/outflow discrepancies detected. Exercise extreme caution.",
        "rec_critical": "🔴 CRITICAL - Wallet is under extremely high risk! Potential illicit flows, mixer activity, or suspicious patterns detected. Suspension or strict monitoring recommended.",
        "footer_text": "Crypto Threat Watch - Cybersecurity Department Special Report.",
        
        # Yangi tranzaksiyalar jadvali uchun
        "tx_history_title": "Full Transaction History (With external wallets)",
        "table_num": "№",
        "table_direction": "Type",
        "table_counterparty": "Counterparty (From/To)",
        "table_amount": "Amount",
        "table_asset": "Asset",
        "no_txs_text": "No transactions with external wallets detected."
    }
}


def _set_cell_background(cell, fill_hex: str) -> None:
    """Jadval yacheykasining fon rangini o'zgartiradi."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150) -> None:
    """Yacheyka ichidagi matn chetlarini (padding) sozlaydi."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_table_borders(table) -> None:
    """Jadval atrofidagi chiziqlarni och-kulrang qiladi."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # Border turlari
    borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    for border_name in borders:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')  # Light Grey
        tblBorders.append(border)
        
    tblPr.append(tblBorders)


def _make_table_header_repeat(table) -> None:
    """Jadval sarlavhasi har sahifada takrorlanishini ta'minlaydi."""
    try:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)
    except Exception:
        pass


def generate_docx_report(
    data: dict,
    risk_level: str,
    risk_emoji: str,
    lang: str,
    output_path: str,
) -> None:
    """Professional Word (.docx) formatidagi hisobotni yaratadi."""
    lang = lang.lower() if lang in REPORT_TEXTS else "uz"
    t = REPORT_TEXTS[lang]

    doc = docx.Document()

    # Sahifa chetlarini sozlash (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Standart stil
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_TEXT

    # 1. HEADER (Sarlavha)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run(t["title"])
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run(t["subtitle"])
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_SECONDARY
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Chiziq ajratgich
    line_p = doc.add_paragraph()
    line_p.paragraph_format.space_after = Pt(12)
    line_run = line_p.add_run("━" * 50)
    line_run.font.color.rgb = COLOR_SECONDARY
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. INFO SECTION (Hamyon ma'lumotlari)
    h_info = doc.add_paragraph()
    h_info_run = h_info.add_run(t["wallet_info"])
    h_info_run.font.size = Pt(14)
    h_info_run.font.bold = True
    h_info_run.font.color.rgb = COLOR_PRIMARY
    h_info.paragraph_format.space_after = Pt(6)

    # Info jadvali
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(info_table)

    _UZT = timezone(timedelta(hours=5))
    labels = [
        (t["network"], data.get("network", "—")),
        (t["address"], data.get("address", "—")),
        (t["status"], data.get("status", "unknown").upper()),
        (t["risk_level"], f"{risk_emoji} {risk_level}"),
        (t["date"], datetime.now(_UZT).strftime("%Y-%m-%d %H:%M (UTC+5)")),
    ]

    # Risk turi bo'yicha rangni aniqlash
    risk_key = "LOW"
    for rk in RISK_COLORS:
        if rk in risk_level.upper():
            risk_key = rk
            break
    risk_hex = RISK_COLORS[risk_key]

    for idx, (label, val) in enumerate(labels):
        row = info_table.rows[idx]
        
        # Label cell
        cell_lbl = row.cells[0]
        cell_lbl.paragraphs[0].add_run(label).bold = True
        _set_cell_background(cell_lbl, "F2F4F7")
        _set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        
        # Val cell
        cell_val = row.cells[1]
        run_val = cell_val.paragraphs[0].add_run(val)
        if label == t["risk_level"]:
            run_val.bold = True
            _set_cell_background(cell_val, risk_hex)
            run_val.font.color.rgb = COLOR_WHITE if risk_key != "MEDIUM" else COLOR_TEXT
        else:
            _set_cell_background(cell_val, "FFFFFF")
            
        _set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. ASSETS TABLE (Aktivlar jadvali)
    h_assets = doc.add_paragraph()
    h_assets_run = h_assets.add_run(t["assets_summary"])
    h_assets_run.font.size = Pt(14)
    h_assets_run.font.bold = True
    h_assets_run.font.color.rgb = COLOR_PRIMARY
    h_assets.paragraph_format.space_after = Pt(6)

    assets = data.get("assets", [])
    if not assets and "total_income" in data:
        assets = [{
            "symbol": data.get("network", "Native"),
            "balance": data.get("current_balance", "—"),
            "income": data.get("total_income", "—"),
            "outcome": data.get("total_outcome", "—"),
            "net": data.get("net_balance", "—"),
            "volume": data.get("total_volume", "—"),
            "tx_count": data.get("tx_count", 0),
        }]

    # Aktiv jadvalini yaratish
    asset_headers = [
        t["table_symbol"], t["table_balance"], t["table_income"], 
        t["table_outcome"], t["table_net"], t["table_volume"], t["table_txs"]
    ]
    
    asset_table = doc.add_table(rows=1 + len(assets), cols=7)
    asset_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(asset_table)
    _make_table_header_repeat(asset_table)

    # Headerlarni yozish
    hdr_cells = asset_table.rows[0].cells
    for i, title in enumerate(asset_headers):
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[i].paragraphs[0].add_run(title)
        run.bold = True
        run.font.color.rgb = COLOR_WHITE
        _set_cell_background(hdr_cells[i], "1F4E79")
        _set_cell_margins(hdr_cells[i], top=100, bottom=100, left=80, right=80)

    # Ma'lumotlarni to'ldirish
    for row_idx, asset in enumerate(assets, 1):
        row_cells = asset_table.rows[row_idx].cells
        bg_color = "F9FBFD" if row_idx % 2 == 0 else "FFFFFF"
        
        vals = [
            asset.get("symbol", "?"),
            asset.get("balance", "—"),
            asset.get("income", "—"),
            asset.get("outcome", "—"),
            asset.get("net", "—"),
            asset.get("volume", "—"),
            str(asset.get("tx_count", 0))
        ]
        
        for col_idx, val in enumerate(vals):
            row_cells[col_idx].paragraphs[0].text = val
            _set_cell_background(row_cells[col_idx], bg_color)
            _set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=80, right=80)
            if col_idx == 0:
                row_cells[col_idx].paragraphs[0].runs[0].bold = True
            else:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 4. SWAPS SECTION (Ichki almashtirishlar)
    h_swaps = doc.add_paragraph()
    h_swaps_run = h_swaps.add_run(t["swaps_title"])
    h_swaps_run.font.size = Pt(14)
    h_swaps_run.font.bold = True
    h_swaps_run.font.color.rgb = COLOR_PRIMARY
    h_swaps.paragraph_format.space_after = Pt(6)

    swaps = data.get("swaps", [])
    if swaps:
        intro_p = doc.add_paragraph()
        intro_p.add_run(t["swaps_intro"])
        intro_p.paragraph_format.space_after = Pt(6)

        swaps_table = doc.add_table(rows=1 + len(swaps), cols=4)
        swaps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(swaps_table)
        _make_table_header_repeat(swaps_table)

        swaps_headers = [t["table_time"], t["table_tx_hash"], t["table_sent"], t["table_received"]]
        hdr_swaps = swaps_table.rows[0].cells
        for i, title in enumerate(swaps_headers):
            hdr_swaps[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hdr_swaps[i].paragraphs[0].add_run(title)
            run.bold = True
            run.font.color.rgb = COLOR_WHITE
            _set_cell_background(hdr_swaps[i], "1F4E79")
            _set_cell_margins(hdr_swaps[i], top=100, bottom=100, left=80, right=80)

        for row_idx, swap in enumerate(swaps, 1):
            row_cells = swaps_table.rows[row_idx].cells
            bg_color = "F9FBFD" if row_idx % 2 == 0 else "FFFFFF"

            dt_str = "—"
            ts = swap.get("timestamp")
            if ts:
                try:
                    dt_str = datetime.fromtimestamp(ts, tz=timezone.utc).strftime("%Y-%m-%d %H:%M")
                except Exception:
                    dt_str = str(ts)

            tx_hash = swap.get("tx_hash", "—")
            short_hash = f"{tx_hash[:10]}...{tx_hash[-8:]}" if len(tx_hash) > 20 else tx_hash

            vals = [dt_str, short_hash, swap.get("from_desc", "—"), swap.get("to_desc", "—")]

            for col_idx, val in enumerate(vals):
                row_cells[col_idx].paragraphs[0].text = val
                _set_cell_background(row_cells[col_idx], bg_color)
                _set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=80, right=80)
                if col_idx in (2, 3):
                    row_cells[col_idx].paragraphs[0].runs[0].bold = True
    else:
        doc.add_paragraph(t["no_swaps"]).paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. YEARLY STATS TABLE (Yillik statistika)
    yearly_stats = data.get("yearly_stats", {})
    if yearly_stats:
        h_yearly = doc.add_paragraph()
        h_yearly_run = h_yearly.add_run(t["yearly_title"])
        h_yearly_run.font.size = Pt(14)
        h_yearly_run.font.bold = True
        h_yearly_run.font.color.rgb = COLOR_PRIMARY
        h_yearly.paragraph_format.space_after = Pt(6)

        native_sym = "GRAM" if data.get("network") == "TON" else ("ETH" if data.get("network") == "ETH" else "TRX")

        yearly_table = doc.add_table(rows=1 + len(yearly_stats), cols=3)
        yearly_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(yearly_table)
        _make_table_header_repeat(yearly_table)

        yearly_headers = [t["yearly_year"], f"{t['table_income']} ({native_sym})", f"{t['table_outcome']} ({native_sym})"]
        hdr_yr = yearly_table.rows[0].cells
        for i, title in enumerate(yearly_headers):
            hdr_yr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hdr_yr[i].paragraphs[0].add_run(title)
            run.bold = True
            run.font.color.rgb = COLOR_WHITE
            _set_cell_background(hdr_yr[i], "1F4E79")
            _set_cell_margins(hdr_yr[i], top=100, bottom=100, left=80, right=80)

        for row_idx, (yr, stats) in enumerate(sorted(yearly_stats.items(), reverse=True), 1):
            row_cells = yearly_table.rows[row_idx].cells
            bg_color = "F9FBFD" if row_idx % 2 == 0 else "FFFFFF"

            row_cells[0].paragraphs[0].text = str(yr)
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].paragraphs[0].text = f"{stats.get('in', 0.0):,.4f}"
            row_cells[2].paragraphs[0].text = f"{stats.get('out', 0.0):,.4f}"

            for col_idx in range(3):
                _set_cell_background(row_cells[col_idx], bg_color)
                _set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=80, right=80)
                if col_idx > 0:
                    row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5.5. FULL TRANSACTION HISTORY (Barcha tranzaksiyalar tarixi jadvali)
    h_txs = doc.add_paragraph()
    h_txs_run = h_txs.add_run(t["tx_history_title"])
    h_txs_run.font.size = Pt(14)
    h_txs_run.font.bold = True
    h_txs_run.font.color.rgb = COLOR_PRIMARY
    h_txs.paragraph_format.space_after = Pt(6)

    normal_transfers = data.get("normal_transfers", [])
    # UTC+5 (Toshkent vaqti)
    UZT = timezone(timedelta(hours=5))
    if normal_transfers:
        # Eng yangi tranzaksiya eng yuqorida (reverse=True)
        sorted_txs = sorted(normal_transfers, key=lambda x: x.get("timestamp", 0), reverse=True)

        txs_table = doc.add_table(rows=1 + len(sorted_txs), cols=6)
        txs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(txs_table)
        _make_table_header_repeat(txs_table)

        txs_headers = [
            t["table_num"], 
            t["table_time"], 
            t["table_direction"], 
            t["table_counterparty"], 
            t["table_amount"], 
            t["table_asset"]
        ]
        
        hdr_txs = txs_table.rows[0].cells
        for i, title in enumerate(txs_headers):
            hdr_txs[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hdr_txs[i].paragraphs[0].add_run(title)
            run.bold = True
            run.font.color.rgb = COLOR_WHITE
            _set_cell_background(hdr_txs[i], "1F4E79")
            _set_cell_margins(hdr_txs[i], top=100, bottom=100, left=80, right=80)

        # Sahifa kengligi ~6.5 dyum. Ustun kengliklari:
        # №: 0.4 in, Vaqt: 1.1 in, Turi: 0.6 in, Hamkor: 2.5 in, Miqdor: 1.1 in, Valyuta: 0.8 in
        col_widths = [Inches(0.4), Inches(1.1), Inches(0.6), Inches(2.5), Inches(1.1), Inches(0.8)]

        for row_idx, tx in enumerate(sorted_txs, 1):
            row_cells = txs_table.rows[row_idx].cells
            bg_color = "F9FBFD" if row_idx % 2 == 0 else "FFFFFF"

            dt_str = "—"
            ts = tx.get("timestamp")
            if ts:
                try:
                    dt_str = datetime.fromtimestamp(ts, tz=UZT).strftime("%Y-%m-%d %H:%M")
                except Exception:
                    dt_str = str(ts)

            direction = tx.get("direction", "in").upper()
            dir_str = "📥 IN" if direction == "IN" else "📤 OUT"

            counterparty = tx.get("counterparty", "—")
            short_cp = counterparty
            if len(counterparty) > 28:
                short_cp = f"{counterparty[:12]}...{counterparty[-12:]}"

            amount_str = f"{tx.get('amount', 0.0):,.4f}"
            asset_str = tx.get("symbol", "?")

            vals = [str(row_idx), dt_str, dir_str, short_cp, amount_str, asset_str]

            for col_idx, val in enumerate(vals):
                cell = row_cells[col_idx]
                p = cell.paragraphs[0]
                run = p.add_run(val)
                _set_cell_background(cell, bg_color)
                _set_cell_margins(cell, top=80, bottom=80, left=60, right=60)
                
                # Tekislash (Alignment)
                if col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif col_idx == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if direction == "IN":
                        run.font.color.rgb = RGBColor(46, 204, 113) # Green
                    else:
                        run.font.color.rgb = RGBColor(231, 76, 60) # Red
                elif col_idx == 4:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif col_idx == 5:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Ustun kengliklarini kiritish
        for row in txs_table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
    else:
        doc.add_paragraph(t["no_txs_text"])

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 6. RISK ASSESSMENT (Xavf darajasi tahlili)
    h_risk = doc.add_paragraph()
    h_risk_run = h_risk.add_run(t["risk_details_title"])
    h_risk_run.font.size = Pt(14)
    h_risk_run.font.bold = True
    h_risk_run.font.color.rgb = COLOR_PRIMARY
    h_risk.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(t["risk_score_desc"]).paragraph_format.space_after = Pt(6)

    # Callout quti yaratish (xavfsizlik tavsiyalari)
    rec_table = doc.add_table(rows=1, cols=1)
    rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(rec_table)
    
    cell_rec = rec_table.rows[0].cells[0]
    _set_cell_background(cell_rec, risk_hex)
    _set_cell_margins(cell_rec, top=120, bottom=120, left=180, right=180)

    # Tavsiya matni
    rec_text = t["rec_low"]
    if "MEDIUM" in risk_level.upper():
        rec_text = t["rec_med"]
    elif "HIGH" in risk_level.upper():
        rec_text = t["rec_high"]
    elif "CRITICAL" in risk_level.upper():
        rec_text = t["rec_critical"]

    p_rec = cell_rec.paragraphs[0]
    run_rec = p_rec.add_run(rec_text)
    run_rec.font.bold = True
    run_rec.font.color.rgb = COLOR_WHITE if risk_key != "MEDIUM" else COLOR_TEXT
    run_rec.font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # 7. FOOTER (Hujjat osti)
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(t["footer_text"])
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = COLOR_SECONDARY
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fayl saqlash
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
